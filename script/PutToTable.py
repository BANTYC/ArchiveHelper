from docx import Document
import xlwt
import os


def export_samus(source_path, export_path):
    first_row = ['№ п/п', 'Полевой шифр', '№ по Описи', 'Коробка', 'Описание', 'Фотографии', 'Рисунок', 'Примечание']
    wb = xlwt.Workbook()
    ws = wb.add_sheet('Лист1')

    for i, cell in enumerate(first_row):
        ws.write(0, i, cell)

    input_path = source_path
    output_path = export_path
    current_row = 1

    files = os.listdir(input_path)

    for file in files:
        doc = Document(os.path.normpath(input_path + '/' + file))
        table = doc.tables[1]

        rows_to_write = []
        rows_to_write.append(table.cell(0, 0).text.replace('Коробка №', '').replace('\n', ''))

        rows_to_write.append([])
        buffer = [a.text for a in table.column_cells(1)]
        rows_to_write[1] = buffer[1:]

        rows_to_write.append([])
        buffer = [a.text for a in table.column_cells(2)]
        rows_to_write[2] = buffer[1:]

        rows_to_write.append([])
        buffer = [a.text for a in table.column_cells(0)]
        rows_to_write[3] = buffer[1:]

        for column1, column2, column3 in zip(rows_to_write[3], rows_to_write[1], rows_to_write[2]):
            ws.write(current_row, 1, column1)
            ws.write(current_row, 3, rows_to_write[0])
            ws.write(current_row, 4, column2)
            ws.write(current_row, 7, column3)
            current_row += 1

        wb.save(output_path + '/SamusData.xls')


def export_passport(source_path, export_path):
    first_row = ['№ п/п', '№ КП', '№ ИК', 'Шифр', 'Тип', 'Год исследования', 'Исследователь',
                 'Наименование коллекции (предмета)', 'Описание коллекции', 'Обстоятельства происхождения',
                 'Археологический памятник', 'Кол-во предметов', 'Размеры', 'Материал, техника', 'Датировка',
                 'Культурная принадлежность', 'Место происхождения коллекции', 'Топография', 'Сохранность']

    input_path = source_path
    output_path = export_path

    wb = xlwt.Workbook()
    ws = wb.add_sheet('Лист1')
    for i, cell in enumerate(first_row):
        ws.write(0, i, cell)

    files = os.listdir(input_path)
    files.sort()

    for i, file in enumerate(files):
        doc = Document(os.path.normpath(input_path + '/' + file))
        table = doc.tables[0]

        column_names = {}
        column_names['№ п/п'] = ''
        column_names['№ КП'] = table.cell(0, 0).text.replace('КП:', '').replace('\n', '')
        column_names['№ ИК'] = table.cell(0, 1).text.replace('ИК:', '').replace('\n', '')
        column_names['Шифр'] = table.cell(0, 2).text.replace('ШИФР:', '').replace('\n', '')
        column_names['Тип'] = 'Археологическая коллекция'
        column_names['Год исследования'] = ''
        column_names['Исследователь'] = table.cell(3, 3).text.replace('Автор:', '').replace('\n', '')
        column_names['Наименование коллекции (предмета)'] = table.cell(1, 3).text.replace('Наименование:', '').replace(
            '\n', '')
        column_names['Описание коллекции'] = table.cell(5, 0).text.replace('Описание:', '').replace('\n', '')
        column_names['Обстоятельства происхождения'] = ''
        column_names['Археологический памятник'] = table.cell(4, 3).text.replace('Место и время создания:', '').replace(
            '\n', '')
        column_names['Кол-во предметов'] = table.cell(0, 3).text.replace('Количество предметов:', '').replace('\n', '')
        column_names['Размеры'] = table.cell(7, 0).text.replace('Размер:', '').replace('\n', '')
        column_names['Материал, техника'] = table.cell(2, 3).text.replace('Материал и техника изготовления:',
                                                                          '').replace('\n', '')
        column_names['Датировка'] = ''
        column_names['Культурная принадлежность'] = 'Не установлено'
        column_names['Место происхождения коллекции'] = ''
        column_names['Топография'] = ''
        column_names['Сохранность'] = table.cell(8, 0).text.replace('Сохранность:', '').replace('\n', '')

        row_to_write = column_names.values()
        for j, cell in enumerate(row_to_write):
            ws.write(i + 1, j, cell)

    wb.save(output_path + '/PassportData.xls')


input_path_samus = os.path.realpath('../input/samus')
input_path_pasport = os.path.realpath('../input/passport')
output_path = os.path.realpath('../output')

export_samus(input_path_samus, output_path)
export_passport(input_path_pasport, output_path)
